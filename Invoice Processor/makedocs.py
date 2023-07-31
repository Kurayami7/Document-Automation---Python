from docx import Document
import random
import math
import os
import openpyxl
from docx import Document

def makeInvoices(numFiles):
    products = ["Parka", "Boots", "Snowshoes", "Climbing Rope", "Oxygen Tank", "Ice Pick", "Crampons"]

    # Invoice loop
    for i in range(numFiles):
        
        # Create Randomized invoice
        invoiceNum = "100" + str(i).zfill(4)
        productList = {}
        for j in range(random.randint(1,10)):
            product = products[random.randint(0,len(products)-1)]
            if product in productList:
                productList[product] += 1
            else:
                productList[product] = 1
        subtot = round(random.random()*10**(random.randint(3, 4)), 2)
        tax = round(subtot*0.13, 2)
        total = round(subtot + tax, 2)

        # Create doc from random invoice
        aDoc = Document()
        aDoc.add_heading("INV" + invoiceNum)
        pProd = aDoc.add_paragraph("PRODUCTS\n")
        for key in productList.keys():
            pProd.add_run(f"{key}:{productList[key]}\n")
        aDoc.add_paragraph(f"SUBTOTAL:{subtot}\nTAX:{tax}\nTOTAL:{total}")
        aDoc.save(f"INV{invoiceNum}.docx")


makeInvoices(200)




# Path to the folder containing the docx files
folder_path = "C:\\Users\\Arief\\Desktop\\Georgian@ILAC\\Semester 2\\Document Automation - Python\\Assignment 2"


def extract_data_from_docx(docx_path):
    document = Document(docx_path)

    # Extract Invoice ID
    invoice_id = document.paragraphs[0].text.strip()

    # Extract products and quantities
    products = {}
    is_products_section = False
    for para in document.paragraphs:
        if para.text.strip().startswith("PRODUCTS"):
            is_products_section = True
        elif para.text.strip().startswith("SUBTOTAL"):
            is_products_section = False

        if is_products_section:
            parts = para.text.split(":")
            if len(parts) == 2:
                product, quantity = parts[0].strip(), parts[1].strip()
                products[product] = int(quantity)
            else:
                product = parts[0].strip()
                products[product] = 0

    # Extract Subtotal, Tax, and Total
    subtotal, tax, total = [float(val.split(":")[1]) for val in document.paragraphs[-1].text.split()]

    # Calculating total quantity by summing the values in the products dictionary
    total_quantity = sum(products.values())

    # print(f"Products: {products}")
    # print(f"Total Quantity: {total_quantity}")
    # print(f"Subtotal: {subtotal}")
    # print(f"Tax: {tax}")
    # print(f"Total: {total}")

    return invoice_id, total_quantity, subtotal, tax, total


def create_excel_output(resultant_data):
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    headers = ["Invoice Number", "Total Quantity", "Subtotal", "Tax", "Total"]
    sheet.append(headers)

    for invoice_data in resultant_data:
        sheet.append(invoice_data)

    return workbook

def main():
    resultant_data = []

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            docx_path = os.path.join(folder_path, filename)
            invoice_id, total_quantity, subtotal, tax, total = extract_data_from_docx(docx_path)

            # total_quantity = sum(products.values())
            row_data = [invoice_id, total_quantity, subtotal, tax, total]
            resultant_data.append(row_data)

    workbook = create_excel_output(resultant_data)
    workbook.save("resultantOutput.xlsx")

if __name__ == "__main__":
    main()